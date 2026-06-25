from docx import Document

from comparison_v3.l4_parser import L4Parser
from comparison_v3.sop_parser import SOPParser


class DocumentExtractor:
    """
    Document Extractor V3

    Responsibilities
    ----------------
    • Detect document type
    • Parse L4 documents
    • Parse SOP documents
    • Return standardized data structure
    """

    def __init__(self):

        self.l4_parser = L4Parser()

        self.sop_parser = SOPParser()

    # ---------------------------------------------------------
    # MAIN ENTRY
    # ---------------------------------------------------------

    def extract(self, file_path):

        filename = str(file_path).lower()

        if "l4" in filename:
            print("USING L4 PARSER")
            return self.extract_l4(file_path)

        print("USING SOP PARSER")
        return self.extract_sop(file_path)

    # ---------------------------------------------------------
    # L4
    # ---------------------------------------------------------

    def extract_l4(self, file_path):

        data = self.l4_parser.parse(file_path)

        data["document_type"] = "L4"

        return data

    # ---------------------------------------------------------
    # SOP
    # ---------------------------------------------------------

    def extract_sop(self, file_path):

        data = self.sop_parser.parse(file_path)

        data["document_type"] = "SOP"

        return data

    # ---------------------------------------------------------
    # DOCUMENT DETECTION
    # ---------------------------------------------------------

    def is_l4_document(self, file_path):

        document = Document(file_path)

        text = []

        # Paragraphs

        for paragraph in document.paragraphs:

            value = paragraph.text.strip()

            if value:

                text.append(

                    value.lower()

                )

        # Tables

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    value = cell.text.strip()

                    if value:

                        text.append(

                            value.lower()

                        )

        content = "\n".join(text)

        indicators = [

            "aris",

            "appendix",

            "process model",

            "model attributes",

            "organization",

            "organization unit",

            "it system",

            "it systems",

            "application system",

            "start event",

            "end event",

            "process step",

            "activity",

            "activities"

        ]

        score = sum(

            1

            for indicator in indicators

            if indicator in content

        )

        return score >= 3

    # ---------------------------------------------------------
    # SUMMARY
    # ---------------------------------------------------------

    def summary(self, data):

        return {

            "document_type":

                data.get(

                    "document_type",

                    ""

                ),

            "process_name":

                data.get(

                    "process_name",

                    ""

                ),

            "activities":

                len(

                    data.get(

                        "activities",

                        []

                    )

                ),

            "roles":

                len(

                    data.get(

                        "roles",

                        []

                    )

                ),

            "applications":

                len(

                    data.get(

                        "applications",

                        []

                    )

                ),

            "controls":

                len(

                    data.get(

                        "controls",

                        []

                    )

                ),

            "risks":

                len(

                    data.get(

                        "risks",

                        []

                    )

                ),

            "inputs":

                len(

                    data.get(

                        "inputs",

                        []

                    )

                ),

            "outputs":

                len(

                    data.get(

                        "outputs",

                        []

                    )

                ),

            "steps":

                len(

                    data.get(

                        "steps",

                        []

                    )

                )

        }


# ---------------------------------------------------------
# TEST
# ---------------------------------------------------------

if __name__ == "__main__":

    extractor = DocumentExtractor()

    try:

        result = extractor.extract(

            "sample.docx"

        )

        print("=" * 80)

        print("DOCUMENT TYPE")

        print(result["document_type"])

        print("=" * 80)

        print()

        print(extractor.summary(result))

    except FileNotFoundError:

        print()

        print("Sample document not found.")

        print("Replace sample.docx with an actual file.")