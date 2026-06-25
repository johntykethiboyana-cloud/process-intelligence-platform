import re
from collections import Counter
from docx import Document
from streamlit import text


class SOPParser:

    def __init__(self):

        self.applications = {
            "sap": "SAP",
            "gcss": "GCSS",
            "aris": "ARIS",
            "sharepoint": "SharePoint",
            "outlook": "Outlook",
            "excel": "Excel",
            "power bi": "Power BI",
            "teams": "Teams",
            "athena": "Athena",
            "salesforce": "Salesforce",
            "servicenow": "ServiceNow",
            "oracle": "Oracle",
            "d365": "Dynamics 365",
            "gpm": "GPM",
            "citrix": "Citrix",
            "mdm": "MDM",
            "cods": "CODS",
            "mci": "MCI"
        }

        self.roles = [
            "manager",
            "analyst",
            "specialist",
            "coordinator",
            "controller",
            "agent",
            "processor",
            "approver",
            "reviewer",
            "owner",
            "requestor",
            "requester",
            "administrator",
            "admin",
            "supervisor",
            "operator"
        ]

        self.action_words = {
            "create",
            "update",
            "delete",
            "review",
            "approve",
            "validate",
            "verify",
            "generate",
            "receive",
            "send",
            "upload",
            "download",
            "assign",
            "submit",
            "cancel",
            "complete",
            "click",
            "select",
            "enter",
            "open",
            "search",
            "navigate",
            "check",
            "maintain",
            "modify",
            "save",
            "capture",
            "populate",
            "choose",
            "view",
            "edit",
            "close",
            "process",
            "confirm",
            "record",
            "register",
            "identify",
            "notify",
            "communicate",
            "draft"
        }

    def clean(self, text):

        if not text:
            return ""

        text = str(text)
        text = text.replace("\n", " ")
        text = re.sub(r"\s+", " ", text)

        return text.strip()

    def unique(self, values):

        result = []
        seen = set()

        for value in values:

            value = self.clean(value)

            if not value:
                continue

            key = value.lower()

            if key not in seen:
                seen.add(key)
                result.append(value)

        return result

    def extract_text(self, file_path):

        doc = Document(file_path)

        lines = []

        for para in doc.paragraphs:

            text = self.clean(para.text)

            if text:
                lines.append(text)

        for table in doc.tables:

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    txt = self.clean(cell.text)

                    if txt:
                        cells.append(txt)

                if cells:
                    for cell_text in cells:
                        lines.append(cell_text)

        return lines

    def extract_process_name(self, lines):

        for line in lines[:20]:

            lower = line.lower()

            if "version" in lower:
                continue

            if len(line) < 8:
                continue

            return line

        return "Unknown Process"

    def extract_role(self, text):

        lower = text.lower()

        role_patterns = {
            "regional head of cpm": "Regional Head of CPM",
            "price setter": "Price Setter",
            "cpm": "CPM",
            "contract product manager": "CPM"
        }

        for pattern, role in role_patterns.items():
            if pattern in lower:
                return role

        return ""

    def extract_application(self, text):

        lower = text.lower()

        for app, name in self.applications.items():

            if app in lower:
                return name

        return ""

    def extract_activity(self, text):

        text = self.clean(text)

        if not text:
            return ""

        lower = text.lower()

        ignore_starts = [
    "purpose",
    "scope",
    "process map",
    "process definition",
    "appendix",
    "requirements",
    "contacts",
    "ownership",
    "index",
    "change log",
    "glossary",
    "communication matrix",
    "review frequency",
    "manage price negotiation for"
]

        # Ignore SOP detailed steps
        if lower.startswith("step "):
            return ""

        # Ignore headers
        for item in ignore_starts:
            if lower.startswith(item):
                return ""

        # Ignore URLs
        if text.startswith("http"):
            return ""

        activity_starts = [
            "review",
            "download",
            "analyse",
            "analyze",
            "prepare",
            "check",
            "update",
            "trigger",
            "send",
            "manage",
            "finalise",
            "finalize",
            "approve",
            "create",
            "receive"
        ]
        for action in activity_starts:
            if action in lower:
                return text

        return ""    

    def parse(self, file_path):    

        lines = self.extract_text(file_path)
        print("\n" + "="*100)
        print("ALL SOP LINES")
        print("="*100)

        for i, line in enumerate(lines):
            print(f"{i+1}: {line}")

        print("="*100)


        result = {
            "process_name": "",
            "activities": [],
            "roles": [],
            "applications": [],
            "controls": [],
            "risks": [],
            "inputs": [],
            "outputs": [],
            "steps": []
        }

        result["process_name"] = self.extract_process_name(lines)

        step_no = 1

        for line in lines:

            lower = line.lower()

            role = self.extract_role(line)

            if role:
                result["roles"].append(role)

            app = self.extract_application(line)

            if app:
                result["applications"].append(app)

            if any(
                word in lower
                for word in [
                    "control",
                    "review",
                    "validate",
                    "verify",
                    "approval",
                    "sign off",
                    "reconcile"
                ]
            ):
                result["controls"].append(line)

            if any(
                word in lower
                for word in [
                    "risk",
                    "error",
                    "failure",
                    "exception",
                    "delay",
                    "incorrect",
                    "duplicate"
                ]
            ):
                result["risks"].append(line)

            if any(
                word in lower
                for word in [
                    "request",
                    "input",
                    "template",
                    "form",
                    "document"
                ]
            ):
                result["inputs"].append(line)

            if any(
                word in lower
                for word in [
                    "output",
                    "report",
                    "generated",
                    "created",
                    "submitted",
                    "completed"
                ]
            ):
                result["outputs"].append(line)

            activity = self.extract_activity(line)

            if activity:

                result["activities"].append(activity)

                result["steps"].append({
                    "step_no": step_no,
                    "activity": activity,
                    "role": role,
                    "application": app
                })

                step_no += 1

        result["activities"] = self.unique(result["activities"])
        result["roles"] = self.unique(result["roles"])
        result["applications"] = self.unique(result["applications"])
        result["controls"] = self.unique(result["controls"])
        result["risks"] = self.unique(result["risks"])
        result["inputs"] = self.unique(result["inputs"])
        result["outputs"] = self.unique(result["outputs"])

        result["summary"] = {
            "process_name": result["process_name"],
            "activity_count": len(result["activities"]),
            "role_count": len(result["roles"]),
            "application_count": len(result["applications"]),
            "control_count": len(result["controls"]),
            "risk_count": len(result["risks"]),
            "input_count": len(result["inputs"]),
            "output_count": len(result["outputs"]),
            "step_count": len(result["steps"])
        }

        print("\n" + "=" * 80)
        print("SOP PARSER DEBUG")
        print("=" * 80)

        print("Process :", result["process_name"])
        print("Activities :", len(result["activities"]))
        print("\nACTIVITIES FOUND")
        print("=" * 80)

        for activity in result["activities"]:
            print(activity)

        print("=" * 80)
        print("Roles :", len(result["roles"]))
        print("Applications :", len(result["applications"]))
        print("Steps :", len(result["steps"]))

        print("=" * 80)

        print("\n" + "=" * 100)
        print("RAW SOP CONTENT")
        print("=" * 100)

        for i, line in enumerate(lines[:300]):
            print(f"{i+1} : {line}")

        print("=" * 100)

        return result


if __name__ == "__main__":

    print("STARTING SOP PARSER")

    parser = SOPParser()

    result = parser.parse(
        r"sample_files\SOP_Manage price negotiation for ocean plus tender 1.docx"
    )

    print("\nRESULT")
    print(result)