from docx import Document
import re


class L4Parser:

    def parse(self, file_path):

        doc = Document(file_path)

        data = {
            "process_name": "",
            "activities": [],
            "roles": [],
            "applications": [],
            "controls": [],
            "risks": [],
            "inputs": [],
            "outputs": [],
            "start_event": "",
            "end_event": "",
            "steps": []
        }

        paragraphs = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)

        # --------------------------------------------------
        # PROCESS NAME
        # --------------------------------------------------

        for line in paragraphs:

            lower = line.lower()

            if (
                "created by" in lower or
                "process manual" in lower or
                "version" in lower or
                "classification" in lower
            ):
                continue

            if len(line) > 10:
                data["process_name"] = line
                break

        # --------------------------------------------------
        # TABLE PARSING
        # --------------------------------------------------

        for table in doc.tables:

            headers = [
                cell.text.strip().lower()
                for cell in table.rows[0].cells
            ]

            activity_col = None
            role_col = None
            input_col = None
            output_col = None
            app_col = None

            for i, h in enumerate(headers):

                if "activity" in h:
                    activity_col = i

                elif "organization" in h or "role" in h:
                    role_col = i

                elif "input" in h:
                    input_col = i

                elif "output" in h:
                    output_col = i

                elif "it" in h or "application" in h or "system" in h:
                    app_col = i

            if activity_col is None:
                continue

            for row in table.rows[1:]:

                cells = row.cells

                activity = ""

                role = ""

                application = ""

                input_value = ""

                output_value = ""

                if activity_col < len(cells):
                    activity = cells[activity_col].text.strip()

                if role_col is not None and role_col < len(cells):
                    role = cells[role_col].text.strip()

                if app_col is not None and app_col < len(cells):
                    application = cells[app_col].text.strip()

                if input_col is not None and input_col < len(cells):
                    input_value = cells[input_col].text.strip()

                if output_col is not None and output_col < len(cells):
                    output_value = cells[output_col].text.strip()

                if not activity:
                    continue

                data["activities"].append(activity)

                if role:
                    data["roles"].append(role)

                if input_value:
                    data["inputs"].append(input_value)

                if output_value:
                    data["outputs"].append(output_value)

                if application:

                    apps = re.split("[;,]", application)

                    for app in apps:

                        app = app.strip()

                        if app:
                            data["applications"].append(app)

                step = {

                    "step_no": len(data["steps"]) + 1,

                    "activity": activity,

                    "role": role,

                    "application": application,

                    "input": input_value,

                    "output": output_value,

                    "description": "",

                    "controls": [],

                    "risks": [],

                    "documents": [],

                    "screenshots": []

                }

                data["steps"].append(step)

        # --------------------------------------------------
        # METADATA
        # --------------------------------------------------

        for line in paragraphs:

            lower = line.lower()

            if lower.startswith("start event"):
                data["start_event"] = line

            elif lower.startswith("end event"):
                data["end_event"] = line

        # --------------------------------------------------
        # APPENDIX PARSING
        # --------------------------------------------------

        current = None

        ignore = {

            "Name",
            "Description",
            "Links",
            "Appendix"

        }

        for line in paragraphs:

            clean = line.strip()

            if clean == "Activities":
                current = "activities"
                continue

            elif clean == "Organizations":
                current = "roles"
                continue

            elif clean == "IT":
                current = "applications"
                continue

            elif clean == "Data":
                current = "inputs"
                continue

            if current is None:
                continue

            if clean in ignore:
                continue

            if len(clean) < 3:
                continue

            data[current].append(clean)

        # --------------------------------------------------
        # STEP ENRICHMENT
        # --------------------------------------------------

        for step in data["steps"]:

            activity = step["activity"].lower()

            for paragraph in paragraphs:

                text = paragraph.strip()

                lower = text.lower()

                if activity in lower:
                    step["description"] += text + "\n"

                if any(

                    word in lower

                    for word in [

                        "approve",

                        "review",

                        "validate",

                        "mandatory",

                        "maker",

                        "checker"

                    ]

                ):

                    step["controls"].append(text)

                if any(

                    word in lower

                    for word in [

                        "risk",

                        "exception",

                        "error",

                        "duplicate"

                    ]

                ):

                    step["risks"].append(text)

                if any(

                    word in lower

                    for word in [

                        "request",

                        "form",

                        "template",

                        "attachment",

                        "instruction"

                    ]

                ):

                    step["documents"].append(text)

                if "figure" in lower or "screenshot" in lower:

                    step["screenshots"].append(text)

        # --------------------------------------------------
        # REMOVE DUPLICATES
        # --------------------------------------------------

        for key in [

            "activities",
            "roles",
            "applications",
            "inputs",
            "outputs"

        ]:

            data[key] = sorted(

                list(

                    set(

                        item.strip()

                        for item in data[key]

                        if item.strip()

                    )

                )

            )

        return data