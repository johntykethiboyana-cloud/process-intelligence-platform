from docx import Document

from comparison.l4_parser import L4Parser
from comparison.sop_parser import SOPParser


class DocumentExtractor:

    def __init__(self):

        self.l4_parser = L4Parser()
        self.sop_parser = SOPParser()

    # -------------------------------------------------------

    def extract(self, file_path):

        if self.is_l4_document(file_path):

            return self.l4_parser.parse(file_path)

        return self.sop_parser.parse(file_path)

    # -------------------------------------------------------

    def is_l4_document(self, file_path):

        """
        Detect whether the uploaded document is
        an ARIS L4 Report.

        Returns:
            True  -> L4 Report
            False -> SOP
        """

        doc = Document(file_path)

        full_text = []

        for p in doc.paragraphs:

            txt = p.text.strip()

            if txt:

                full_text.append(txt.lower())

        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    txt = cell.text.strip()

                    if txt:

                        full_text.append(txt.lower())

        text = "\n".join(full_text)

        indicators = [

            "aris",

            "process model",

            "activities",

            "organizations",

            "it systems",

            "start event",

            "end event",

            "model attributes",

            "appendix"

        ]

        score = 0

        for word in indicators:

            if word in text:

                score += 1

        return score >= 3