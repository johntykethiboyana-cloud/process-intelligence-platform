import re
from docx import Document


class SOPParser:

    def parse(self, file_path):

        document = Document(file_path)

        data = {

            "process_name": "",

            "activities": [],

            "roles": [],

            "applications": [],

            "controls": [],

            "risks": [],

            "inputs": [],

            "outputs": [],

            "start_event": "",

            "end_event": "",

            "steps": []

        }

        paragraphs = []

        for p in document.paragraphs:

            text = p.text.strip()

            if text:

                paragraphs.append(text)

        # --------------------------------------------------
        # PROCESS NAME
        # --------------------------------------------------

        for line in paragraphs:

            if line.lower().startswith("classification"):

                continue

            if len(line) > 10:

                data["process_name"] = line

                break

        # --------------------------------------------------
        # ACTIVITIES
        # --------------------------------------------------

        activity_pattern = re.compile(r"^\d+(\.\d+)+\s+")

        current = None

        description = []

        for line in paragraphs:

            if activity_pattern.match(line):

                if current:

                    current["description"] = "\n".join(description)

                    data["steps"].append(current)

                activity = activity_pattern.sub("", line).strip()

                current = {

                    "activity": activity,

                    "description": ""

                }

                data["activities"].append(activity)

                description = []

                continue

            if current:

                description.append(line)

        if current:

            current["description"] = "\n".join(description)

            data["steps"].append(current)

        # --------------------------------------------------
        # TABLES
        # --------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    text = cell.text.strip()

                    if not text:

                        continue

                    lower = text.lower()

                    if "request" in lower:

                        data["inputs"].append(text)

                    if any(

                        word in lower

                        for word in [

                            "completed",

                            "updated",

                            "created",

                            "submitted",

                            "finalised"

                        ]

                    ):

                        data["outputs"].append(text)

                    if any(

                        word in lower

                        for word in [

                            "approve",

                            "review",

                            "validate",

                            "mandatory",

                            "checker",

                            "maker"

                        ]

                    ):

                        data["controls"].append(text)

                    if any(

                        word in lower

                        for word in [

                            "risk",

                            "error",

                            "duplicate",

                            "incorrect",

                            "delay",

                            "exception"

                        ]

                    ):

                        data["risks"].append(text)

        # --------------------------------------------------
        # PARAGRAPHS
        # --------------------------------------------------

        applications = [

            "Athena",

            "Excel",

            "Outlook",

            "SAP",

            "Power BI",

            "SharePoint",

            "FASWeb",

            "Citrix",

            "ServiceNow",

            "ARIS"

        ]

        roles = [

            "CPM",

            "Contract Product Manager",

            "Price Setter",

            "Requester",

            "Approver",

            "Reviewer",

            "Operations",

            "Treasury",

            "Finance"

        ]

        for line in paragraphs:

            lower = line.lower()

            for app in applications:

                if app.lower() in lower:

                    data["applications"].append(app)

            for role in roles:

                if role.lower() in lower:

                    data["roles"].append(role)

            if "start" in lower and not data["start_event"]:

                data["start_event"] = line

            if "end" in lower and not data["end_event"]:

                data["end_event"] = line

            if "request" in lower:

                data["inputs"].append(line)

            if any(

                word in lower

                for word in [

                    "completed",

                    "updated",

                    "submitted",

                    "created",

                    "finalised"

                ]

            ):

                data["outputs"].append(line)

            if any(

                word in lower

                for word in [

                    "approve",

                    "review",

                    "validate",

                    "mandatory",

                    "checker",

                    "maker"

                ]

            ):

                data["controls"].append(line)

            if any(

                word in lower

                for word in [

                    "risk",

                    "duplicate",

                    "incorrect",

                    "delay",

                    "exception"

                ]

            ):

                data["risks"].append(line)

        # --------------------------------------------------
        # CLEAN
        # --------------------------------------------------

        for key in [

            "activities",

            "roles",

            "applications",

            "controls",

            "risks",

            "inputs",

            "outputs"

        ]:

            data[key] = sorted(

                list(

                    set(

                        x.strip()

                        for x in data[key]

                        if x.strip()

                    )

                )

            )

        return data