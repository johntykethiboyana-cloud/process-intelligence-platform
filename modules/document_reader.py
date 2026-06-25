from docx import Document

class DocumentReader:

    def __init__(self, file_path):
        self.file_path = file_path
        self.document = Document(file_path)

    def get_paragraphs(self):
        paragraphs = []

        for para in self.document.paragraphs:
            text = para.text.strip()

            if text != "":
                paragraphs.append(text)

        return paragraphs

    def get_tables(self):
        return self.document.tables

    def get_paragraph_details(self):
        details = []

        for para in self.document.paragraphs:
            if para.text.strip() != "":
                details.append({
                    "text": para.text.strip(),
                    "style": para.style.name
                })

        return details