import re
from collections import Counter
from docx import Document


class SOPParser:

    """
    SOP Parser V2

    Part 1
    ------------------------
    • Initialisation
    • Utility Functions
    • Text Cleaning
    • Heading Detection
    • Step Detection
    • Action Extraction
    • Business Object Extraction
    • Keyword Extraction
    • Application Normalization
    """

    def __init__(self):

        self.applications = {

            "sap": "SAP",
            "gcss": "GCSS",
            "sharepoint": "SharePoint",
            "aris": "ARIS",
            "excel": "Excel",
            "outlook": "Outlook",
            "power bi": "Power BI",
            "service now": "ServiceNow",
            "servicenow": "ServiceNow",
            "athena": "Athena",
            "citrix": "Citrix",
            "salesforce": "Salesforce",
            "teams": "Teams"

        }

        self.actions = [

            "create",
            "update",
            "delete",
            "approve",
            "review",
            "validate",
            "verify",
            "select",
            "click",
            "enter",
            "open",
            "close",
            "search",
            "submit",
            "assign",
            "upload",
            "download",
            "receive",
            "send",
            "copy",
            "move",
            "generate",
            "cancel",
            "link",
            "unlink"

        ]

        self.stop_words = {

            "the",
            "a",
            "an",
            "of",
            "for",
            "to",
            "and",
            "or",
            "with",
            "by",
            "in",
            "on",
            "this",
            "that",
            "is",
            "are",
            "be",
            "using",
            "please",
            "kindly"

        }

    # --------------------------------------------------
    # CLEAN TEXT
    # --------------------------------------------------

    def clean_text(self, text):

        if not text:
            return ""

        text = str(text)

        text = text.replace("\n", " ")

        text = re.sub(r"\s+", " ", text)

        text = text.strip()

        return text

    # --------------------------------------------------
    # HEADING
    # --------------------------------------------------

    def is_heading(self, text):

        text = self.clean_text(text)

        if len(text) < 3:
            return False

        if len(text) < 70 and text.isupper():
            return True

        if re.match(r"^\d+(\.\d+)*\s+[A-Z]", text):
            return True

        return False

    # --------------------------------------------------
    # STEP
    # --------------------------------------------------

    def is_step(self, text):

        patterns = [

            r"^\d+\.",

            r"^\d+\)",

            r"^\d+\.\d+",

            r"^Step\s+\d+",

            r"^\(\d+\)",

            r"^\d+\s+-"

        ]

        return any(

            re.match(

                pattern,

                text,

                re.IGNORECASE

            )

            for pattern in patterns

        )

    # --------------------------------------------------
    # ACTION
    # --------------------------------------------------

    def extract_action(self, text):

        lower = self.clean_text(text).lower()

        for action in self.actions:

            if re.search(

                rf"\b{re.escape(action)}\b",

                lower

            ):

                return action

        return ""

    # --------------------------------------------------
    # BUSINESS OBJECT
    # --------------------------------------------------

    def extract_business_object(self, text):

        action = self.extract_action(text)

        if not action:

            return self.clean_text(text)

        match = re.search(

            rf"{re.escape(action)}\s+(.*)",

            text,

            re.IGNORECASE

        )

        if match:

            return match.group(1).strip()

        return self.clean_text(text)

    # --------------------------------------------------
    # APPLICATION
    # --------------------------------------------------

    def normalize_application(self, text):

        lower = self.clean_text(text).lower()

        for key, value in self.applications.items():

            if key in lower:

                return value

        return ""

    # --------------------------------------------------
    # KEYWORDS
    # --------------------------------------------------

    def extract_keywords(self, text):

        words = re.findall(

            r"[A-Za-z0-9]+",

            self.clean_text(text).lower()

        )

        words = [

            word

            for word in words

            if len(word) > 2

            and word not in self.stop_words

        ]

        return sorted(

            list(

                Counter(words)

            )

        )

    # --------------------------------------------------
    # UNIQUE
    # --------------------------------------------------

    def unique(self, values):

        output = []

        seen = set()

        for value in values:

            value = self.clean_text(value)

            if not value:

                continue

            key = value.lower()

            if key not in seen:

                seen.add(key)

                output.append(value)

        return output

    # --------------------------------------------------
    # STEP OBJECT
    # --------------------------------------------------

    def build_step(

        self,

        step_no,

        activity,

        description=""

    ):

        return {

            "step_no": step_no,

            "activity": activity,

            "clean_activity": self.clean_text(activity),

            "action": self.extract_action(activity),

            "business_object": self.extract_business_object(activity),

            "keywords": self.extract_keywords(activity),

            "description": description,

            "role": "",

            "application": "",

            "controls": [],

            "risks": [],

            "documents": [],

            "screenshots": [],

            "metadata": {}

        }
    # --------------------------------------------------
    # INITIAL DATA
    # --------------------------------------------------

    def initialise_data(self):

        return {

            "process_name": "",

            "activities": [],

            "roles": [],

            "applications": [],

            "controls": [],

            "risks": [],

            "inputs": [],

            "outputs": [],

            "start_event": "",

            "end_event": "",

            "steps": []

        }

    # --------------------------------------------------
    # READ PARAGRAPHS
    # --------------------------------------------------

    def read_paragraphs(self, document):

        paragraphs = []

        for paragraph in document.paragraphs:

            text = self.clean_text(

                paragraph.text

            )

            if text:

                paragraphs.append(text)

        return paragraphs

    # --------------------------------------------------
    # PROCESS NAME
    # --------------------------------------------------

    def extract_process_name(self, paragraphs):

        ignore = [

            "classification",

            "version",

            "author",

            "created by",

            "effective date",

            "page",

            "document"

        ]

        for line in paragraphs:

            lower = line.lower()

            if any(

                word in lower

                for word in ignore

            ):

                continue

            if len(line) > 10:

                return line

        return "Unknown Process"

    # --------------------------------------------------
    # PARSE STEPS
    # --------------------------------------------------

    def parse_steps(self, paragraphs, data):

        current_step = None

        description = []

        for line in paragraphs:

            if self.is_step(line):

                if current_step:

                    current_step["description"] = "\n".join(

                        description

                    )

                    data["steps"].append(

                        current_step

                    )

                activity = re.sub(

                    r"^\d+(\.\d+)*",

                    "",

                    line

                )

                activity = activity.strip(". ")

                current_step = self.build_step(

                    len(data["steps"]) + 1,

                    activity

                )

                data["activities"].append(

                    activity

                )

                description = []

                continue

            if current_step:

                description.append(line)

        if current_step:

            current_step["description"] = "\n".join(

                description

            )

            data["steps"].append(

                current_step

            )

    # --------------------------------------------------
    # TABLE PARSER
    # --------------------------------------------------

    def parse_tables(

        self,

        document,

        data

    ):

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    text = self.clean_text(

                        cell.text

                    )

                    if not text:

                        continue

                    lower = text.lower()

                    application = self.normalize_application(

                        text

                    )

                    if application:

                        data["applications"].append(

                            application

                        )

                    if any(

                        word in lower

                        for word in [

                            "request",

                            "form",

                            "input"

                        ]

                    ):

                        data["inputs"].append(

                            text

                        )

                    if any(

                        word in lower

                        for word in [

                            "completed",

                            "created",

                            "updated",

                            "submitted",

                            "generated",

                            "output"

                        ]

                    ):

                        data["outputs"].append(

                            text

                        )

                    if any(

                        word in lower

                        for word in [

                            "approve",

                            "review",

                            "validate",

                            "mandatory",

                            "checker",

                            "maker"

                        ]

                    ):

                        data["controls"].append(

                            text

                        )

                    if any(

                        word in lower

                        for word in [

                            "risk",

                            "error",

                            "duplicate",

                            "delay",

                            "incorrect",

                            "exception"

                        ]

                    ):

                        data["risks"].append(

                            text

                        )

    # --------------------------------------------------
    # ROLE EXTRACTION
    # --------------------------------------------------

    def extract_roles(

        self,

        paragraphs,

        data

    ):

        roles = [

            "CPM",

            "Contract Product Manager",

            "Requester",

            "Approver",

            "Reviewer",

            "Treasury",

            "Finance",

            "Operations",

            "Price Setter"

        ]

        for line in paragraphs:

            lower = line.lower()

            for role in roles:

                if role.lower() in lower:

                    data["roles"].append(

                        role

                    )

    # --------------------------------------------------
    # APPLICATION EXTRACTION
    # --------------------------------------------------

    def extract_applications(

        self,

        paragraphs,

        data

    ):

        for line in paragraphs:

            app = self.normalize_application(

                line

            )

            if app:

                data["applications"].append(

                    app

                )

    # --------------------------------------------------
    # EVENTS
    # --------------------------------------------------

    def extract_events(

        self,

        paragraphs,

        data

    ):

        for line in paragraphs:

            lower = line.lower()

            if (

                "start"

                in lower

                and

                not data["start_event"]

            ):

                data["start_event"] = line

            if (

                "end"

                in lower

                and

                not data["end_event"]

            ):

                data["end_event"] = line
    # --------------------------------------------------
    # ENRICH STEPS
    # --------------------------------------------------

    def enrich_steps(self, paragraphs, data):

        control_words = [

            "approve",
            "approval",
            "review",
            "validate",
            "verification",
            "mandatory",
            "maker",
            "checker",
            "authorization"

        ]

        risk_words = [

            "risk",
            "error",
            "exception",
            "duplicate",
            "incorrect",
            "delay",
            "failure"

        ]

        document_words = [

            "template",
            "form",
            "attachment",
            "document",
            "request",
            "guideline",
            "instruction"

        ]

        screenshot_words = [

            "figure",
            "screenshot",
            "screen shot",
            "screen print",
            "image"

        ]

        for step in data["steps"]:

            descriptions = []

            controls = []

            risks = []

            documents = []

            screenshots = []

            activity = step["activity"].lower()

            for paragraph in paragraphs:

                text = self.clean_text(paragraph)

                lower = text.lower()

                if activity in lower:

                    descriptions.append(text)

                if any(

                    word in lower

                    for word in control_words

                ):

                    controls.append(text)

                if any(

                    word in lower

                    for word in risk_words

                ):

                    risks.append(text)

                if any(

                    word in lower

                    for word in document_words

                ):

                    documents.append(text)

                if any(

                    word in lower

                    for word in screenshot_words

                ):

                    screenshots.append(text)

            step["description"] = "\n".join(

                self.unique(descriptions)

            )

            step["controls"] = self.unique(

                controls

            )

            step["risks"] = self.unique(

                risks

            )

            step["documents"] = self.unique(

                documents

            )

            step["screenshots"] = self.unique(

                screenshots

            )

            data["controls"].extend(

                step["controls"]

            )

            data["risks"].extend(

                step["risks"]

            )

    # --------------------------------------------------
    # INPUTS / OUTPUTS
    # --------------------------------------------------

    def extract_inputs_outputs(

        self,

        paragraphs,

        data

    ):

        input_words = [

            "request",

            "input",

            "received",

            "receive",

            "template",

            "form"

        ]

        output_words = [

            "completed",

            "updated",

            "submitted",

            "created",

            "generated",

            "output",

            "closed",

            "approved"

        ]

        for paragraph in paragraphs:

            lower = paragraph.lower()

            if any(

                word in lower

                for word in input_words

            ):

                data["inputs"].append(

                    paragraph

                )

            if any(

                word in lower

                for word in output_words

            ):

                data["outputs"].append(

                    paragraph

                )

    # --------------------------------------------------
    # STEP METADATA
    # --------------------------------------------------

    def build_metadata(

        self,

        data

    ):

        for index, step in enumerate(

            data["steps"],

            start=1

        ):

            step["metadata"] = {

                "step_no": index,

                "action": step["action"],

                "business_object": step["business_object"],

                "keyword_count": len(

                    step["keywords"]

                ),

                "application": step["application"],

                "role": step["role"]

            }

    # --------------------------------------------------
    # CLEANUP
    # --------------------------------------------------

    def cleanup(

        self,

        data

    ):

        for key in [

            "activities",

            "roles",

            "applications",

            "controls",

            "risks",

            "inputs",

            "outputs"

        ]:

            data[key] = self.unique(

                data[key]

            )

        return data
    # --------------------------------------------------
    # VALIDATION
    # --------------------------------------------------

    def validate(self, data):

        if not data["process_name"]:

            data["process_name"] = "Unknown Process"

        for step in data["steps"]:

            step["activity"] = self.clean_text(
                step["activity"]
            )

            step["clean_activity"] = self.clean_text(
                step["activity"]
            )

            step["action"] = self.extract_action(
                step["activity"]
            )

            step["business_object"] = self.extract_business_object(
                step["activity"]
            )

            step["keywords"] = self.extract_keywords(
                step["activity"]
            )

            if step["application"]:

                step["application"] = self.normalize_application(
                    step["application"]
                )

        return data

    # --------------------------------------------------
    # SUMMARY
    # --------------------------------------------------

    def build_summary(self, data):

        data["summary"] = {

            "activities": len(data["activities"]),

            "roles": len(data["roles"]),

            "applications": len(data["applications"]),

            "controls": len(data["controls"]),

            "risks": len(data["risks"]),

            "inputs": len(data["inputs"]),

            "outputs": len(data["outputs"]),

            "steps": len(data["steps"])

        }

        return data

    # --------------------------------------------------
    # MAIN PARSER
    # --------------------------------------------------

    def parse(self, file_path):

        document = Document(file_path)

        data = self.initialise_data()

        paragraphs = self.read_paragraphs(document)

        # ------------------------------------------

        data["process_name"] = self.extract_process_name(
            paragraphs
        )

        # ------------------------------------------

        self.parse_steps(
            paragraphs,
            data
        )

        # ------------------------------------------

        self.parse_tables(
            document,
            data
        )

        # ------------------------------------------

        self.extract_roles(
            paragraphs,
            data
        )

        # ------------------------------------------

        self.extract_applications(
            paragraphs,
            data
        )

        # ------------------------------------------

        self.extract_events(
            paragraphs,
            data
        )

        # ------------------------------------------

        self.extract_inputs_outputs(
            paragraphs,
            data
        )

        # ------------------------------------------

        self.enrich_steps(
            paragraphs,
            data
        )

        # ------------------------------------------

        self.build_metadata(
            data
        )

        # ------------------------------------------

        data = self.cleanup(
            data
        )

        # ------------------------------------------

        data = self.validate(
            data
        )

        # ------------------------------------------

        data = self.build_summary(
            data
        )

        return data
