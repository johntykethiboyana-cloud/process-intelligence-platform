from docx import Document

from comparison_v2.l4_parser import L4Parser
from comparison_v2.sop_parser import SOPParser
from comparison_v2.activity_extractor import ActivityExtractor


class DocumentExtractor:
    """
    Document Extractor V2

    Responsibilities
    ----------------
    1. Detect document type
    2. Parse L4 document
    3. Parse SOP document
    4. Enrich SOP with business activities
    """

    def __init__(self):

        self.l4_parser = L4Parser()
        self.sop_parser = SOPParser()
        self.activity_extractor = ActivityExtractor()

    # -------------------------------------------------------
    # MAIN ENTRY
    # -------------------------------------------------------

    def extract(self, file_path):
        is_l4 = self.is_l4_document(file_path)

        print("\n============================")
        print("FILE:", file_path)
        print("IS_L4:", is_l4)
        print("============================")

        if is_l4:
            print("USING L4 PARSER")
            return self.extract_l4(file_path)

        print("USING SOP PARSER")
        return self.extract_sop(file_path)

    # -------------------------------------------------------
    # L4 PARSER
    # -------------------------------------------------------

    def extract_l4(self, file_path):

        return self.l4_parser.parse(file_path)

    # -------------------------------------------------------
    # SOP PARSER
    # -------------------------------------------------------

    def extract_sop(self, file_path):

        sop_data = self.sop_parser.parse(file_path)

        print("\n========== SOP PARSER OUTPUT ==========")
        print(type(sop_data))
        print(sop_data)
        print("=======================================")

        return self.enrich_sop(sop_data)

    # -------------------------------------------------------
    # DOCUMENT TYPE DETECTION
    # -------------------------------------------------------

    def is_l4_document(self, file_path):

        doc = Document(file_path)

        text = []

        for paragraph in doc.paragraphs:

            if paragraph.text.strip():

                text.append(paragraph.text.lower())

        for table in doc.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():

                        text.append(cell.text.lower())

        content = "\n".join(text)

        indicators = [

            "aris",
            "process model",
            "activities",
            "organizations",
            "it systems",
            "start event",
            "end event",
            "model attributes",
            "appendix"

        ]

        score = sum(
        1
            for indicator in indicators
            if indicator in content
)

        print("\n========== DOCUMENT CHECK ==========")
        print("FILE:", file_path)
        print("SCORE:", score)

        for indicator in indicators:
            if indicator in content:
                print("FOUND:", indicator)

        print("===================================")
    
        # TEMPORARY TEST
        if "aris" in content:
            return True
        if "model attributes" in content:
            return True
        if "it systems" in content and "organizations" in content:
            return True

        return False
    
    # -------------------------------------------------------
    # SOP ENRICHMENT
    # -------------------------------------------------------

    def enrich_sop(self, sop_data):

        business_activities = []
        business_steps = []

        # ---------------------------------------------------
        # Collect all activities from parser + steps
        # ---------------------------------------------------

        candidates = []

        candidates.extend(

            sop_data.get("activities", [])

        )

        for step in sop_data.get("steps", []):

            if isinstance(step, dict):

                text = (
                    step.get("activity")
                    or step.get("step")
                    or step.get("text")
                    or ""
                )

            else:

                text = str(step)

            if text.strip():

                candidates.append(text)

        candidates = self.unique(candidates)

        # ---------------------------------------------------
        # Enrich activities
        # ---------------------------------------------------

        for activity in candidates:

            sections = {

                "Activities": [

                    activity

                ]

            }

            extracted = self.activity_extractor.extract(

                sections

            )

            if extracted:

                for item in extracted:

                    if item["type"] in [

                        "ACTIVITY",
                        "APPROVAL",
                        "VALIDATION",
                        "COMMUNICATION"

                    ]:

                        business_activities.append(

                            item["text"]

                        )

                        business_steps.append({

                            "activity": item["text"],

                            "original_activity": activity,

                            "type": item["type"],

                            "heading": item.get(

                                "heading", ""

                            )

                        })

            else:

                business_activities.append(activity)

                business_steps.append({

                    "activity": activity,

                    "original_activity": activity,

                    "type": "ACTIVITY",

                    "heading": ""

                })

        # ---------------------------------------------------
        # Save enriched data
        # ---------------------------------------------------

        sop_data["activities"] = self.unique(candidates)

        sop_data["business_activities"] = self.unique(

            business_activities

        )

        sop_data["business_steps"] = business_steps

        return self.build_summary(sop_data)

    # -------------------------------------------------------
    # UNIQUE VALUES
    # -------------------------------------------------------

    def unique(self, values):

        output = []

        seen = set()

        for value in values:

            value = str(value).strip()

            if not value:

                continue

            key = value.lower()

            if key not in seen:

                seen.add(key)

                output.append(value)

        return output

    # -------------------------------------------------------
    # BUILD SUMMARY
    # -------------------------------------------------------

    def build_summary(self, sop_data):

        sop_data["summary"] = {

            "activities": len(

                sop_data.get("activities", [])

            ),

            "business_activities": len(

                sop_data.get("business_activities", [])

            ),

            "steps": len(

                sop_data.get("steps", [])

            ),

            "business_steps": len(

                sop_data.get("business_steps", [])

            ),

            "applications": len(

                sop_data.get("applications", [])

            ),

            "roles": len(

                sop_data.get("roles", [])

            ),

            "controls": len(

                sop_data.get("controls", [])

            ),

            "risks": len(

                sop_data.get("risks", [])

            ),

            "inputs": len(

                sop_data.get("inputs", [])

            ),

            "outputs": len(

                sop_data.get("outputs", [])

            )

        }

        return sop_data