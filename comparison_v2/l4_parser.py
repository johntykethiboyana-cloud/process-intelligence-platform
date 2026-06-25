from docx import Document
import re
from collections import Counter


class L4Parser:

    """
    L4 Parser Version 2

    Part 1
    --------
    - Initialisation
    - Utility Functions
    - Text Cleaning
    - Action Extraction
    - Business Object Extraction
    - Application Normalization
    - Keyword Extraction
    """

    def __init__(self):

        self.application_map = {

            "gcss": "GCSS",
            "gpm": "GPM",
            "sap": "SAP",
            "aris": "ARIS",
            "sharepoint": "SharePoint",
            "outlook": "Outlook",
            "excel": "Excel",
            "power bi": "Power BI",
            "salesforce": "Salesforce",
            "teams": "Teams"

        }

        self.action_words = [

            "create",
            "update",
            "delete",
            "approve",
            "review",
            "validate",
            "verify",
            "submit",
            "open",
            "close",
            "select",
            "click",
            "enter",
            "assign",
            "search",
            "upload",
            "download",
            "link",
            "unlink",
            "copy",
            "move",
            "generate",
            "receive",
            "send",
            "cancel",
            "complete"

        ]

        self.stop_words = {

            "the",
            "a",
            "an",
            "of",
            "for",
            "to",
            "in",
            "on",
            "with",
            "by",
            "and",
            "or",
            "is",
            "are",
            "be",
            "this",
            "that",
            "from",
            "into",
            "using",
            "user"

        }

    # --------------------------------------------------
    # CLEAN TEXT
    # --------------------------------------------------

    def clean_text(self, text):

        if not text:
            return ""

        text = str(text)

        text = text.replace("\n", " ")

        text = re.sub(r"\s+", " ", text)

        text = re.sub(r"\.{2,}", ".", text)

        text = text.strip()

        return text

    # --------------------------------------------------
    # NORMALIZE APPLICATION
    # --------------------------------------------------

    def normalize_application(self, value):

        value = self.clean_text(value)

        if not value:
            return ""

        lower = value.lower()

        for key, app in self.application_map.items():

            if key in lower:

                return app

        return value

    # --------------------------------------------------
    # EXTRACT ACTION
    # --------------------------------------------------

    def extract_action(self, text):

        text = self.clean_text(text).lower()

        for action in self.action_words:

            if re.search(

                rf"\b{re.escape(action)}\b",

                text

            ):

                return action

        return ""

    # --------------------------------------------------
    # BUSINESS OBJECT
    # --------------------------------------------------

    def extract_business_object(self, text):

        text = self.clean_text(text)

        action = self.extract_action(text)

        if not action:

            return text

        pattern = rf"{re.escape(action)}\s+(.*)"

        match = re.search(

            pattern,

            text,

            re.IGNORECASE

        )

        if match:

            obj = match.group(1)

            obj = obj.strip()

            return obj

        return text

    # --------------------------------------------------
    # KEYWORDS
    # --------------------------------------------------

    def extract_keywords(self, text):

        text = self.clean_text(text).lower()

        words = re.findall(

            r"[a-zA-Z0-9]+",

            text

        )

        words = [

            word

            for word in words

            if len(word) > 2

            and word not in self.stop_words

        ]

        return sorted(

            list(

                Counter(words)

            )

        )

    # --------------------------------------------------
    # REMOVE DUPLICATES
    # --------------------------------------------------

    def unique(self, values):

        result = []

        seen = set()

        for item in values:

            item = self.clean_text(item)

            if not item:

                continue

            key = item.lower()

            if key not in seen:

                seen.add(key)

                result.append(item)

        return result

    # --------------------------------------------------
    # SAFE CELL VALUE
    # --------------------------------------------------

    def get_cell(self, cells, index):

        if index is None:

            return ""

        if index >= len(cells):

            return ""

        return self.clean_text(

            cells[index].text

        )

    # --------------------------------------------------
    # CREATE STEP OBJECT
    # --------------------------------------------------

    def build_step(

        self,

        step_no,

        activity,

        role,

        application,

        input_value,

        output_value

    ):

        return {

            "step_no": step_no,

            "activity": activity,

            "clean_activity": self.clean_text(activity),

            "action": self.extract_action(activity),

            "business_object": self.extract_business_object(activity),

            "keywords": self.extract_keywords(activity),

            "role": role,

            "application": self.normalize_application(application),

            "input": input_value,

            "output": output_value,

            "description": "",

            "controls": [],

            "risks": [],

            "documents": [],

            "screenshots": [],

            "metadata": {}

        }
    # --------------------------------------------------
    # PROCESS NAME
    # --------------------------------------------------

    def extract_process_name(self, paragraphs):

        ignore = [
            "created by",
            "process manual",
            "version",
            "classification",
            "document owner",
            "approved by",
            "effective date",
            "page "
        ]

        for line in paragraphs:

            text = self.clean_text(line)

            if len(text) < 10:
                continue

            lower = text.lower()

            if any(x in lower for x in ignore):
                continue

            return text

        return ""

    # --------------------------------------------------
    # TABLE HEADER DETECTION
    # --------------------------------------------------

    def detect_columns(self, headers):

        columns = {
            "activity": None,
            "role": None,
            "application": None,
            "input": None,
            "output": None
        }

        for index, header in enumerate(headers):

            h = header.lower()

            if any(x in h for x in [
                "activity",
                "function",
                "task",
                "process step"
            ]):
                columns["activity"] = index

            elif any(x in h for x in [
                "organization",
                "role",
                "owner",
                "performer"
            ]):
                columns["role"] = index

            elif any(x in h for x in [
                "application",
                "system",
                "it",
                "tool"
            ]):
                columns["application"] = index

            elif "input" in h:
                columns["input"] = index

            elif "output" in h:
                columns["output"] = index

        return columns

    # --------------------------------------------------
    # READ PARAGRAPHS
    # --------------------------------------------------

    def read_paragraphs(self, doc):

        paragraphs = []

        for paragraph in doc.paragraphs:

            text = self.clean_text(paragraph.text)

            if text:

                paragraphs.append(text)

        return paragraphs

    # --------------------------------------------------
    # READ TABLES
    # --------------------------------------------------

    def parse_tables(self, doc, data):

        for table_number, table in enumerate(doc.tables, start=1):

            if len(table.rows) < 2:
                continue

            headers = [
                self.clean_text(cell.text)
                for cell in table.rows[0].cells
            ]

            mapping = self.detect_columns(headers)

            if mapping["activity"] is None:
                continue

            for row_number, row in enumerate(table.rows[1:], start=2):

                cells = row.cells

                activity = self.get_cell(
                    cells,
                    mapping["activity"]
                )

                if not activity:
                    continue

                role = self.get_cell(
                    cells,
                    mapping["role"]
                )

                application = self.get_cell(
                    cells,
                    mapping["application"]
                )

                input_value = self.get_cell(
                    cells,
                    mapping["input"]
                )

                output_value = self.get_cell(
                    cells,
                    mapping["output"]
                )

                if application:

                    application = self.normalize_application(
                        application
                    )

                step = self.build_step(

                    len(data["steps"]) + 1,

                    activity,

                    role,

                    application,

                    input_value,

                    output_value

                )

                step["metadata"] = {

                    "table": table_number,
                    "row": row_number

                }

                data["steps"].append(step)

                data["activities"].append(activity)

                if role:
                    data["roles"].append(role)

                if application:
                    data["applications"].append(application)

                if input_value:
                    data["inputs"].append(input_value)

                if output_value:
                    data["outputs"].append(output_value)

    # --------------------------------------------------
    # INITIAL PARSE
    # --------------------------------------------------

    def initialise_data(self):

        return {

            "process_name": "",

            "activities": [],

            "roles": [],

            "applications": [],

            "controls": [],

            "risks": [],

            "inputs": [],

            "outputs": [],

            "start_event": "",

            "end_event": "",

            "steps": []

        }
    # --------------------------------------------------
    # APPENDIX PARSER
    # --------------------------------------------------

    def parse_appendix(self, paragraphs, data):

        current = None

        sections = {
            "activities": ["activities", "activity"],
            "roles": ["organizations", "roles"],
            "applications": ["it", "applications", "systems"],
            "inputs": ["data", "documents", "inputs"],
            "outputs": ["outputs"]
        }

        ignore = {
            "name",
            "description",
            "links",
            "appendix",
            "object type",
            "attribute"
        }

        for line in paragraphs:

            text = self.clean_text(line)

            lower = text.lower()

            found = False

            for key, values in sections.items():

                if lower in values:

                    current = key

                    found = True

                    break

            if found:
                continue

            if current is None:
                continue

            if lower in ignore:
                continue

            if len(text) < 3:
                continue

            data[current].append(text)

    # --------------------------------------------------
    # STEP ENRICHMENT
    # --------------------------------------------------

    def enrich_steps(self, paragraphs, data):

        control_words = [

            "approve",
            "approval",
            "review",
            "validate",
            "verification",
            "mandatory",
            "maker",
            "checker",
            "authorization"

        ]

        risk_words = [

            "risk",
            "exception",
            "error",
            "duplicate",
            "failure",
            "incorrect",
            "missing"

        ]

        document_words = [

            "template",
            "form",
            "attachment",
            "document",
            "request",
            "instruction",
            "guideline"

        ]

        screenshot_words = [

            "figure",
            "screen shot",
            "screenshot",
            "screen print",
            "image"

        ]

        for step in data["steps"]:

            activity = step["activity"].lower()

            descriptions = []

            controls = []

            risks = []

            documents = []

            screenshots = []

            for paragraph in paragraphs:

                text = self.clean_text(paragraph)

                lower = text.lower()

                if activity in lower:

                    descriptions.append(text)

                if any(

                    word in lower

                    for word in control_words

                ):

                    controls.append(text)

                if any(

                    word in lower

                    for word in risk_words

                ):

                    risks.append(text)

                if any(

                    word in lower

                    for word in document_words

                ):

                    documents.append(text)

                if any(

                    word in lower

                    for word in screenshot_words

                ):

                    screenshots.append(text)

            step["description"] = "\n".join(

                self.unique(descriptions)

            )

            step["controls"] = self.unique(

                controls

            )

            step["risks"] = self.unique(

                risks

            )

            step["documents"] = self.unique(

                documents

            )

            step["screenshots"] = self.unique(

                screenshots

            )

            data["controls"].extend(

                step["controls"]

            )

            data["risks"].extend(

                step["risks"]

            )

    # --------------------------------------------------
    # EVENTS
    # --------------------------------------------------

    def extract_events(self, paragraphs, data):

        for line in paragraphs:

            lower = line.lower()

            if lower.startswith("start event"):

                data["start_event"] = line

            elif lower.startswith("end event"):

                data["end_event"] = line

    # --------------------------------------------------
    # FINAL CLEANUP
    # --------------------------------------------------

    def cleanup(self, data):

        for key in [

            "activities",

            "roles",

            "applications",

            "controls",

            "risks",

            "inputs",

            "outputs"

        ]:

            data[key] = self.unique(

                data[key]

            )

        return data
    # --------------------------------------------------
    # VALIDATION
    # --------------------------------------------------

    def validate(self, data):

        if not data["process_name"]:

            data["process_name"] = "Unknown Process"

        for step in data["steps"]:

            step["activity"] = self.clean_text(

                step["activity"]

            )

            step["clean_activity"] = self.clean_text(

                step["activity"]

            )

            step["action"] = self.extract_action(

                step["activity"]

            )

            step["business_object"] = self.extract_business_object(

                step["activity"]

            )

            step["keywords"] = self.extract_keywords(

                step["activity"]

            )

            step["application"] = self.normalize_application(

                step["application"]

            )

        return data

    # --------------------------------------------------
    # SUMMARY
    # --------------------------------------------------

    def build_summary(self, data):

        data["summary"] = {

            "activities": len(data["activities"]),

            "roles": len(data["roles"]),

            "applications": len(data["applications"]),

            "controls": len(data["controls"]),

            "risks": len(data["risks"]),

            "inputs": len(data["inputs"]),

            "outputs": len(data["outputs"]),

            "steps": len(data["steps"])

        }

        return data

    # --------------------------------------------------
    # MAIN PARSER
    # --------------------------------------------------

    def parse(self, file_path):

        document = Document(file_path)

        data = self.initialise_data()

        paragraphs = self.read_paragraphs(

            document

        )

        # ----------------------------------------------

        data["process_name"] = self.extract_process_name(

            paragraphs

        )

        # ----------------------------------------------

        self.parse_tables(

            document,

            data

        )

        # ----------------------------------------------

        self.extract_events(

            paragraphs,

            data

        )

        # ----------------------------------------------

        self.parse_appendix(

            paragraphs,

            data

        )

        # ----------------------------------------------

        self.enrich_steps(

            paragraphs,

            data

        )

        # ----------------------------------------------

        data = self.cleanup(

            data

        )

        # ----------------------------------------------

        data = self.validate(

            data

        )

        # ----------------------------------------------

        data = self.build_summary(

            data

        )

        return data
